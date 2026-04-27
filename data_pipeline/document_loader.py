"""
document_loader.py
──────────────────
Unified document ingestion layer.
Supports: PDF, HTML, plain text, Markdown, DOCX, JSON/JSONL, URLs, SQL databases.

Design decisions:
- Dataclass `Document` is the canonical unit across the entire pipeline.
- Each loader is a thin adapter; heavy lifting is delegated to proven libs.
- Async-capable: async def load() wraps sync operations with asyncio.to_thread.
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import re
import sqlite3
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Generator, List, Optional
from urllib.parse import urlparse

import requests

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Canonical data model
# ---------------------------------------------------------------------------
@dataclass
class Document:
    """Represents a single document or document-section entering the pipeline."""
    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.sha256(
                (self.content + json.dumps(self.metadata, sort_keys=True)).encode()
            ).hexdigest()[:16]

    def __len__(self) -> int:
        return len(self.content)

    @property
    def word_count(self) -> int:
        return len(self.content.split())


# ---------------------------------------------------------------------------
# Base loader interface
# ---------------------------------------------------------------------------
class BaseLoader:
    """All loaders implement load() → List[Document]."""

    def load(self) -> List[Document]:
        raise NotImplementedError

    async def aload(self) -> List[Document]:
        return await asyncio.to_thread(self.load)

    @staticmethod
    def _clean_text(text: str) -> str:
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


# ---------------------------------------------------------------------------
# PDF Loader
# ---------------------------------------------------------------------------
class PDFLoader(BaseLoader):
    """
    Loads PDFs page-by-page using pypdf.
    Preserves page numbers in metadata for citation traceability.
    """

    def __init__(self, file_path: str, extract_images: bool = False):
        self.file_path = Path(file_path)
        self.extract_images = extract_images

    def load(self) -> List[Document]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Install pypdf: pip install pypdf")

        reader = PdfReader(str(self.file_path))
        docs: List[Document] = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = self._clean_text(text)
            if len(text) < 20:          # skip nearly-empty pages
                continue
            docs.append(Document(
                content=text,
                metadata={
                    "source": str(self.file_path),
                    "page": page_num,
                    "total_pages": len(reader.pages),
                    "file_type": "pdf",
                    "ingested_at": time.time(),
                }
            ))

        logger.info("PDFLoader: loaded %d pages from %s", len(docs), self.file_path)
        return docs


# ---------------------------------------------------------------------------
# HTML Loader
# ---------------------------------------------------------------------------
class HTMLLoader(BaseLoader):
    """
    Loads HTML from a file path or URL.
    Strips navigation/scripts; preserves semantic structure.
    """

    def __init__(self, source: str, remove_tags: Optional[List[str]] = None):
        self.source = source
        self.remove_tags = remove_tags or ["script", "style", "nav", "footer", "header"]
        self._is_url = source.startswith(("http://", "https://"))

    def load(self) -> List[Document]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("Install beautifulsoup4: pip install beautifulsoup4 lxml")

        if self._is_url:
            response = requests.get(self.source, timeout=30)
            response.raise_for_status()
            html_content = response.text
            source_meta = self.source
        else:
            html_content = Path(self.source).read_text(encoding="utf-8")
            source_meta = self.source

        soup = BeautifulSoup(html_content, "lxml")
        for tag in self.remove_tags:
            for element in soup.find_all(tag):
                element.decompose()

        text = soup.get_text(separator="\n")
        text = self._clean_text(text)
        title = soup.title.string if soup.title else ""

        return [Document(
            content=text,
            metadata={
                "source": source_meta,
                "title": title,
                "file_type": "html",
                "ingested_at": time.time(),
            }
        )]


# ---------------------------------------------------------------------------
# Plain Text / Markdown Loader
# ---------------------------------------------------------------------------
class TextLoader(BaseLoader):
    """Loads .txt or .md files."""

    def __init__(self, file_path: str, encoding: str = "utf-8"):
        self.file_path = Path(file_path)
        self.encoding = encoding

    def load(self) -> List[Document]:
        content = self.file_path.read_text(encoding=self.encoding)
        content = self._clean_text(content)
        return [Document(
            content=content,
            metadata={
                "source": str(self.file_path),
                "file_type": self.file_path.suffix.lstrip("."),
                "ingested_at": time.time(),
            }
        )]


# ---------------------------------------------------------------------------
# DOCX Loader
# ---------------------------------------------------------------------------
class DocxLoader(BaseLoader):
    """Microsoft Word document loader."""

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)

    def load(self) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = DocxDocument(str(self.file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        content = self._clean_text(content)

        return [Document(
            content=content,
            metadata={
                "source": str(self.file_path),
                "file_type": "docx",
                "ingested_at": time.time(),
            }
        )]


# ---------------------------------------------------------------------------
# JSON / JSONL Loader
# ---------------------------------------------------------------------------
class JSONLoader(BaseLoader):
    """
    Loads JSON / JSONL files.
    `content_key` specifies which field holds the main text.
    All other fields are folded into metadata.
    """

    def __init__(self, file_path: str, content_key: str = "text",
                 metadata_keys: Optional[List[str]] = None):
        self.file_path = Path(file_path)
        self.content_key = content_key
        self.metadata_keys = metadata_keys or []

    def _record_to_doc(self, record: dict) -> Optional[Document]:
        content = record.get(self.content_key, "")
        if not content:
            return None
        meta = {k: record[k] for k in self.metadata_keys if k in record}
        meta.update({"source": str(self.file_path), "file_type": "json",
                     "ingested_at": time.time()})
        return Document(content=str(content), metadata=meta)

    def load(self) -> List[Document]:
        docs: List[Document] = []
        suffix = self.file_path.suffix.lower()

        with self.file_path.open(encoding="utf-8") as f:
            if suffix == ".jsonl":
                for line in f:
                    line = line.strip()
                    if line:
                        doc = self._record_to_doc(json.loads(line))
                        if doc:
                            docs.append(doc)
            else:
                data = json.load(f)
                if isinstance(data, list):
                    for record in data:
                        doc = self._record_to_doc(record)
                        if doc:
                            docs.append(doc)
                else:
                    doc = self._record_to_doc(data)
                    if doc:
                        docs.append(doc)

        logger.info("JSONLoader: loaded %d records from %s", len(docs), self.file_path)
        return docs


# ---------------------------------------------------------------------------
# SQL / Database Loader
# ---------------------------------------------------------------------------
class SQLLoader(BaseLoader):
    """
    Pulls rows from a SQLite DB (extend for PostgreSQL via psycopg2).
    Each row becomes a Document; columns are concatenated as key:value pairs.
    """

    def __init__(self, db_path: str, query: str,
                 content_columns: Optional[List[str]] = None):
        self.db_path = db_path
        self.query = query
        self.content_columns = content_columns

    def load(self) -> List[Document]:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.execute(self.query)
        rows = cursor.fetchall()
        conn.close()

        docs: List[Document] = []
        for row in rows:
            row_dict = dict(row)
            if self.content_columns:
                content_parts = [f"{k}: {row_dict[k]}" for k in self.content_columns
                                  if k in row_dict]
            else:
                content_parts = [f"{k}: {v}" for k, v in row_dict.items()]
            content = " | ".join(content_parts)
            docs.append(Document(
                content=content,
                metadata={"source": self.db_path, "file_type": "sql",
                           "ingested_at": time.time(), **row_dict}
            ))

        logger.info("SQLLoader: loaded %d rows", len(docs))
        return docs


# ---------------------------------------------------------------------------
# Directory Loader (recursive multi-format)
# ---------------------------------------------------------------------------
class DirectoryLoader(BaseLoader):
    """
    Recursively walks a directory and dispatches to the correct loader
    based on file extension.

    Supported: .pdf, .html, .htm, .txt, .md, .docx, .json, .jsonl
    """

    EXTENSION_MAP = {
        ".pdf": PDFLoader,
        ".html": HTMLLoader,
        ".htm": HTMLLoader,
        ".txt": TextLoader,
        ".md": TextLoader,
        ".docx": DocxLoader,
        ".json": JSONLoader,
        ".jsonl": JSONLoader,
    }

    def __init__(self, directory: str, glob: str = "**/*",
                 exclude_patterns: Optional[List[str]] = None,
                 max_files: Optional[int] = None):
        self.directory = Path(directory)
        self.glob = glob
        self.exclude_patterns = exclude_patterns or []
        self.max_files = max_files

    def _should_exclude(self, path: Path) -> bool:
        return any(re.search(pat, str(path)) for pat in self.exclude_patterns)

    def load(self) -> List[Document]:
        all_docs: List[Document] = []
        files = list(self.directory.glob(self.glob))

        if self.max_files:
            files = files[:self.max_files]

        for file_path in files:
            if not file_path.is_file():
                continue
            if self._should_exclude(file_path):
                continue
            loader_cls = self.EXTENSION_MAP.get(file_path.suffix.lower())
            if loader_cls is None:
                logger.debug("Skipping unsupported file: %s", file_path)
                continue
            try:
                loader = loader_cls(str(file_path))
                docs = loader.load()
                all_docs.extend(docs)
            except Exception as exc:
                logger.warning("Failed to load %s: %s", file_path, exc)

        logger.info("DirectoryLoader: loaded %d documents from %s",
                    len(all_docs), self.directory)
        return all_docs


# ---------------------------------------------------------------------------
# Unified façade
# ---------------------------------------------------------------------------
class DocumentLoader:
    """
    Entry-point façade. Auto-detects format or delegates by explicit type.

    Usage:
        loader = DocumentLoader()
        docs = loader.load("./docs/annual_report.pdf")
        docs = loader.load("https://example.com/article")
        docs = loader.load("./data/", loader_type="directory")
    """

    def load(self, source: str, loader_type: Optional[str] = None,
             **kwargs) -> List[Document]:
        if loader_type is None:
            loader_type = self._detect_type(source)

        loader_map = {
            "pdf":       PDFLoader,
            "html":      HTMLLoader,
            "url":       HTMLLoader,
            "txt":       TextLoader,
            "md":        TextLoader,
            "docx":      DocxLoader,
            "json":      JSONLoader,
            "jsonl":     JSONLoader,
            "sql":       SQLLoader,
            "directory": DirectoryLoader,
        }

        loader_cls = loader_map.get(loader_type)
        if loader_cls is None:
            raise ValueError(f"Unsupported loader type: {loader_type}")

        loader = loader_cls(source, **kwargs)
        return loader.load()

    async def aload(self, source: str, **kwargs) -> List[Document]:
        return await asyncio.to_thread(self.load, source, **kwargs)

    @staticmethod
    def _detect_type(source: str) -> str:
        if source.startswith(("http://", "https://")):
            return "url"
        p = Path(source)
        if p.is_dir():
            return "directory"
        return p.suffix.lstrip(".").lower() or "txt"
